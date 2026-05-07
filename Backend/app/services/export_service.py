import io
import re
import sys
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Any

from motor.motor_asyncio import AsyncIOMotorDatabase

from app.schemas.analytics import MeetingAnalyticsResponse
from app.schemas.readout import ReadoutResponse
from app.services.analytics_service import AnalyticsService
from app.services.meeting_service import MeetingService
from app.services.readout_service import ReadoutService


@dataclass
class ExportPayload:
    filename: str
    media_type: str
    content: bytes


class ExportService:
    def __init__(self) -> None:
        self.meetings = MeetingService()
        self.readouts = ReadoutService()
        self.analytics = AnalyticsService()

    async def export_meeting(
        self,
        *,
        db: AsyncIOMotorDatabase,
        owner: dict,
        meeting_id: str,
        export_format: str,
    ) -> ExportPayload:
        meeting = await self.meetings.get_meeting(db=db, owner=owner, meeting_id=meeting_id)
        if not meeting:
            raise ValueError("Meeting not found")
        markdown = self._meeting_markdown(meeting)
        title = meeting.get("title") or "Meeting"
        filename_base = self._slugify(title) or "meeting"
        return self._build_payload(
            export_format=export_format,
            filename_base=filename_base,
            title=title,
            markdown=markdown,
        )

    async def export_readout(
        self,
        *,
        db: AsyncIOMotorDatabase,
        owner: dict,
        readout_id: str,
        export_format: str,
    ) -> ExportPayload:
        documents = await db["readouts"].find({"id": readout_id, "owner_id": owner["id"]}).limit(1).to_list(length=1)
        if not documents:
            raise ValueError("Readout not found")
        readout = ReadoutResponse.model_validate(documents[0])
        markdown = self._readout_markdown(readout)
        filename_base = self._slugify(readout.title) or "readout"
        return self._build_payload(
            export_format=export_format,
            filename_base=filename_base,
            title=readout.title,
            markdown=markdown,
        )

    async def export_analytics(
        self,
        *,
        db: AsyncIOMotorDatabase,
        owner: dict,
        export_format: str,
    ) -> ExportPayload:
        analytics = await self.analytics.get_meeting_analytics(db=db, owner=owner)
        markdown = self._analytics_markdown(analytics)
        return self._build_payload(
            export_format=export_format,
            filename_base="notable-analytics",
            title="Notable analytics",
            markdown=markdown,
        )

    def _build_payload(
        self,
        *,
        export_format: str,
        filename_base: str,
        title: str,
        markdown: str,
    ) -> ExportPayload:
        normalized_format = export_format.strip().lower()
        if normalized_format == "markdown":
            return ExportPayload(
                filename=f"{filename_base}.md",
                media_type="text/markdown; charset=utf-8",
                content=markdown.encode("utf-8"),
            )
        if normalized_format == "docx":
            return ExportPayload(
                filename=f"{filename_base}.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=self._build_docx(title=title, markdown=markdown),
            )
        if normalized_format == "pdf":
            return ExportPayload(
                filename=f"{filename_base}.pdf",
                media_type="application/pdf",
                content=self._build_pdf(title=title, markdown=markdown),
            )
        raise ValueError("Unsupported export format")

    def _meeting_markdown(self, meeting: dict) -> str:
        transcript_chunks = meeting.get("transcript_chunks") or []
        action_items = meeting.get("action_items") or []
        participants = [item for item in (meeting.get("participants") or []) if str(item).strip()]
        lines = [
            f"# {meeting.get('title') or 'Meeting'}",
            "",
            "## Snapshot",
            f"- Status: {meeting.get('status') or 'unknown'}",
            f"- Provider: {meeting.get('provider') or 'other'}",
            f"- Template: {meeting.get('summary_template') or 'office_meeting'}",
            f"- Participants: {', '.join(participants) if participants else 'None listed'}",
        ]
        if meeting.get("scheduled_start"):
            lines.append(f"- Scheduled start: {meeting['scheduled_start']}")
        if meeting.get("scheduled_end"):
            lines.append(f"- Scheduled end: {meeting['scheduled_end']}")

        lines.extend(["", "## Summary", "", meeting.get("summary") or "No summary generated yet.", ""])
        lines.extend(["## Action Items", ""])
        if action_items:
            lines.extend([f"- {item}" for item in action_items])
        else:
            lines.append("- No action items captured.")

        lines.extend(["", "## Notes", "", meeting.get("notes_markdown") or "No notes yet.", "", "## Transcript", ""])
        if transcript_chunks:
            for chunk in transcript_chunks:
                speaker = chunk.get("speaker_label") or "Speaker"
                lines.append(f"- **{speaker}:** {chunk.get('transcript_text') or ''}")
        else:
            lines.append("- No transcript captured.")
        return "\n".join(lines).strip() + "\n"

    def _readout_markdown(self, readout: ReadoutResponse) -> str:
        lines = [
            f"# {readout.title}",
            "",
            "## Snapshot",
            f"- Timeframe: {readout.timeframe}",
            f"- Sources: {', '.join(readout.sources)}",
            f"- Created at: {readout.created_at.isoformat()}",
            "",
            "## Summary",
            "",
            readout.summary,
            "",
            "## Key Points",
            "",
        ]
        lines.extend([f"- {item}" for item in readout.key_points] or ["- None captured."])
        lines.extend(["", "## Action Items", ""])
        lines.extend([f"- {item}" for item in readout.action_items] or ["- None captured."])
        lines.extend(["", "## Suggested Replies", ""])
        lines.extend([f"- {item}" for item in readout.suggested_replies] or ["- None captured."])
        if readout.notices:
            lines.extend(["", "## Notices", ""])
            lines.extend([f"- {item}" for item in readout.notices])
        return "\n".join(lines).strip() + "\n"

    def _analytics_markdown(self, analytics: MeetingAnalyticsResponse) -> str:
        lines = [
            "# Notable analytics",
            "",
            "## Overview",
            f"- Total meetings: {analytics.overview.total_meetings}",
            f"- Summarized meetings: {analytics.overview.summarized_meetings}",
            f"- Meetings with recordings: {analytics.overview.meetings_with_recordings}",
            f"- Shared meetings: {analytics.overview.shared_meetings}",
            f"- Total action items: {analytics.overview.total_action_items}",
            f"- Total questions: {analytics.overview.total_questions}",
            f"- Average duration minutes: {analytics.overview.average_duration_minutes}",
            f"- Average participants: {analytics.overview.average_participants}",
            "",
            "## Highlights",
            "",
        ]
        lines.extend([f"- **{item.title}:** {item.body}" for item in analytics.highlights] or ["- No highlights yet."])
        lines.extend(["", "## Provider Breakdown", ""])
        for item in analytics.provider_breakdown:
            lines.append(
                f"- {item.label}: {item.meetings} meetings, avg {item.average_duration_minutes} min, shares {item.share_count}"
            )
        if not analytics.provider_breakdown:
            lines.append("- No provider data yet.")
        lines.extend(["", "## Sharing Breakdown", ""])
        for item in analytics.visibility_breakdown:
            lines.append(f"- {item.label}: {item.meetings} meetings, {item.total_views} views")
        if not analytics.visibility_breakdown:
            lines.append("- No sharing data yet.")
        lines.extend(["", "## Monthly Activity", ""])
        for item in analytics.monthly_activity:
            lines.append(
                f"- {item.label}: {item.meetings} meetings, {item.action_items} action items, {item.words} words"
            )
        if not analytics.monthly_activity:
            lines.append("- No monthly activity yet.")
        lines.extend(["", "## Top Meetings", ""])
        for item in analytics.top_meetings:
            lines.append(
                f"- {item.title}: {item.action_items} actions, {item.questions} questions, {item.duration_minutes} min, {item.share_views} views"
            )
        if not analytics.top_meetings:
            lines.append("- No meetings ranked yet.")
        return "\n".join(lines).strip() + "\n"

    def _ensure_optional_export_packages(self) -> None:
        candidate = (
            Path.home()
            / ".cache"
            / "codex-runtimes"
            / "codex-primary-runtime"
            / "dependencies"
            / "python"
            / "Lib"
            / "site-packages"
        )
        if candidate.exists() and str(candidate) not in sys.path:
            sys.path.insert(0, str(candidate))

    def _build_docx(self, *, title: str, markdown: str) -> bytes:
        self._ensure_optional_export_packages()
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise ValueError("DOCX export dependencies are not available on this runtime") from exc

        document = Document()
        document.add_heading(title, level=0)
        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                document.add_paragraph("")
                continue
            if stripped.startswith("# "):
                document.add_heading(stripped[2:].strip(), level=1)
            elif stripped.startswith("## "):
                document.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("- "):
                document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(self._strip_markdown_inline(stripped))
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _build_pdf(self, *, title: str, markdown: str) -> bytes:
        self._ensure_optional_export_packages()
        try:
            from reportlab.lib import colors  # type: ignore
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
            from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer  # type: ignore
        except ImportError as exc:
            raise ValueError("PDF export dependencies are not available on this runtime") from exc

        background = colors.HexColor("#0b0b0c")
        foreground = colors.HexColor("#f4f4f5")
        muted = colors.HexColor("#a1a1aa")
        border = colors.HexColor("#262626")

        buffer = io.BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=40,
            rightMargin=40,
            topMargin=48,
            bottomMargin=42,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "NotableTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=30,
            textColor=foreground,
            alignment=TA_LEFT,
            spaceAfter=12,
        )
        heading_style = ParagraphStyle(
            "NotableHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=foreground,
            spaceBefore=10,
            spaceAfter=8,
        )
        body_style = ParagraphStyle(
            "NotableBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=muted,
            spaceAfter=6,
        )
        story: list[Any] = [Paragraph(escape(title), title_style), Spacer(1, 6)]

        bullet_buffer: list[ListItem] = []

        def flush_bullets() -> None:
            nonlocal bullet_buffer
            if bullet_buffer:
                story.append(
                    ListFlowable(
                        bullet_buffer,
                        bulletType="bullet",
                        start="circle",
                        bulletFontName="Helvetica",
                        bulletColor=foreground,
                        leftIndent=14,
                    )
                )
                story.append(Spacer(1, 8))
                bullet_buffer = []

        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                flush_bullets()
                continue
            if stripped.startswith("# "):
                flush_bullets()
                story.append(Paragraph(escape(self._strip_markdown_inline(stripped[2:])), title_style))
                continue
            if stripped.startswith("## "):
                flush_bullets()
                story.append(Paragraph(escape(self._strip_markdown_inline(stripped[3:])), heading_style))
                continue
            if stripped.startswith("- "):
                bullet_buffer.append(ListItem(Paragraph(escape(self._strip_markdown_inline(stripped[2:])), body_style)))
                continue
            flush_bullets()
            story.append(Paragraph(escape(self._strip_markdown_inline(stripped)), body_style))

        flush_bullets()

        def paint_theme(canvas, _doc) -> None:
            width, height = A4
            canvas.saveState()
            canvas.setFillColor(background)
            canvas.rect(0, 0, width, height, fill=1, stroke=0)
            canvas.setStrokeColor(border)
            canvas.line(40, height - 34, width - 40, height - 34)
            canvas.setFillColor(foreground)
            canvas.setFont("Helvetica-Bold", 10)
            canvas.drawString(40, height - 24, "Notable export")
            canvas.restoreState()

        document.build(story, onFirstPage=paint_theme, onLaterPages=paint_theme)
        return buffer.getvalue()

    def _strip_markdown_inline(self, value: str) -> str:
        cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
        cleaned = re.sub(r"\*(.*?)\*", r"\1", cleaned)
        cleaned = re.sub(r"`(.*?)`", r"\1", cleaned)
        return cleaned

    def _slugify(self, value: str) -> str:
        parts = "".join(char.lower() if char.isalnum() else "-" for char in value)
        return re.sub(r"-{2,}", "-", parts).strip("-")
